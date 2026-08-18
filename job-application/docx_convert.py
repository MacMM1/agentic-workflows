#!/usr/bin/env python3
"""Convert between .docx and an indexed plain-text format.

to-md:   docx -> .md, one line per paragraph, each tagged with a stable
         paragraph index (and style) so edits can be mapped back later.
to-docx: takes an ORIGINAL docx (for structure/formatting) plus an
         edited .md produced by to-md, and rewrites the text of any
         paragraph whose content changed. Each paragraph is matched by
         index, and the first run's formatting is kept while later runs
         in that paragraph are cleared -- this is built for light,
         word/phrase-level edits, not restructuring (added/removed
         paragraphs will not map correctly).

Usage:
  python job-application/docx_convert.py to-md --input saved-applications/cv-source-materials/cv-master.docx --output .tmp/cv-master.md
  python job-application/docx_convert.py to-docx --original saved-applications/cv-source-materials/cv-master.docx --edited saved-applications/foo/cv-foo.md --output saved-applications/foo/cv-foo.docx
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document

LINE_TAG = re.compile(r"^<!--\s*p(\d+)\s+style=(\S+)\s*-->(.*)$")


def _escape(text: str) -> str:
    """Paragraphs can contain soft line breaks (w:br -> '\\n' in python-docx's
    para.text). Escape them so each paragraph stays on exactly one physical
    line in the .md file -- the to-docx parser relies on that 1:1 mapping."""
    return text.replace("\\", "\\\\").replace("\n", "\\n").replace("\t", "\\t")


def _unescape(text: str) -> str:
    out = []
    i = 0
    while i < len(text):
        if text[i] == "\\" and i + 1 < len(text):
            nxt = text[i + 1]
            if nxt == "n":
                out.append("\n"); i += 2; continue
            if nxt == "t":
                out.append("\t"); i += 2; continue
            if nxt == "\\":
                out.append("\\"); i += 2; continue
        out.append(text[i]); i += 1
    return "".join(out)


def to_md(input_path: Path, output_path: Path) -> None:
    doc = Document(str(input_path))
    lines = []
    for i, para in enumerate(doc.paragraphs):
        style = (para.style.name if para.style else "Normal").replace(" ", "_")
        lines.append(f"<!-- p{i} style={style} -->{_escape(para.text)}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {len(lines)} paragraphs to {output_path}")


def _set_paragraph_text(para, new_text: str) -> None:
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def to_docx(original_path: Path, edited_path: Path, output_path: Path) -> None:
    doc = Document(str(original_path))
    edited_lines = edited_path.read_text(encoding="utf-8").splitlines()

    changed = 0
    seen = set()
    for line in edited_lines:
        m = LINE_TAG.match(line)
        if not m:
            if line.strip():
                print(f"WARNING: skipping line with no p<N> tag: {line[:60]!r}", file=sys.stderr)
            continue
        idx, new_text = int(m.group(1)), _unescape(m.group(3))
        seen.add(idx)
        if idx >= len(doc.paragraphs):
            print(f"WARNING: p{idx} not found in original (paragraph added?), skipping", file=sys.stderr)
            continue
        para = doc.paragraphs[idx]
        if para.text != new_text:
            _set_paragraph_text(para, new_text)
            changed += 1

    missing = set(range(len(doc.paragraphs))) - seen
    if missing:
        print(f"WARNING: edited file is missing {len(missing)} paragraph(s) present in the original "
              f"(indices {sorted(missing)[:10]}{'...' if len(missing) > 10 else ''}) -- they were left unchanged",
              file=sys.stderr)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Wrote {output_path} ({changed} paragraph(s) changed)")


def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    sub = parser.add_subparsers(dest="mode", required=True)

    p1 = sub.add_parser("to-md", help="Extract a .docx into indexed text")
    p1.add_argument("--input", required=True, type=Path)
    p1.add_argument("--output", required=True, type=Path)

    p2 = sub.add_parser("to-docx", help="Apply edited indexed text back onto the original .docx")
    p2.add_argument("--original", required=True, type=Path, help="Unmodified source .docx (structure/formatting base)")
    p2.add_argument("--edited", required=True, type=Path, help="Edited .md produced by to-md")
    p2.add_argument("--output", required=True, type=Path)

    args = parser.parse_args()
    if args.mode == "to-md":
        if not args.input.exists():
            sys.exit(f"Input not found: {args.input}")
        to_md(args.input, args.output)
    else:
        if not args.original.exists():
            sys.exit(f"Original not found: {args.original}")
        if not args.edited.exists():
            sys.exit(f"Edited file not found: {args.edited}")
        to_docx(args.original, args.edited, args.output)


if __name__ == "__main__":
    main()
